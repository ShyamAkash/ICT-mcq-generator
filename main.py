from google import genai
from docx import Document
import json
from UnicodeToLegacy import ConvertToLegacy
import tkinter as tk
from tkinter import ttk

root = tk.Tk()
root.title("Question Generator")
root.geometry("450x300")
root.configure(bg="#1e1e2f")  # dark background

style = ttk.Style()
style.theme_use("clam")

style.configure("TLabel", background="#1e1e2f", foreground="white", font=("Segoe UI", 11))
style.configure("TButton", font=("Segoe UI", 11), padding=6)
style.configure("TRadiobutton", background="#1e1e2f", foreground="white", font=("Segoe UI", 10))
style.configure("TEntry", padding=5)

title = ttk.Label(root, text="ictfromabc Question Generator", font=("Segoe UI", 16, "bold"))
title.pack(pady=15)

frame = tk.Frame(root, bg="#2a2a40", padx=15, pady=15)
frame.pack(pady=10, padx=20, fill="both", expand=True)

ttk.Label(frame, text="Enter Question Topic:").pack(anchor="w")

topic_entry = ttk.Entry(frame, width=40)
topic_entry.pack(pady=5, fill="x")

ttk.Label(frame, text="Select Question Type:").pack(anchor="w", pady=(10, 5))

question_type = tk.StringVar(value="Normal")

ttk.Radiobutton(frame, text="Normal", variable=question_type, value="normal").pack(anchor="w")
ttk.Radiobutton(frame, text="Code", variable=question_type, value="code").pack(anchor="w")
ttk.Radiobutton(frame, text="Statement", variable=question_type, value="statement").pack(anchor="w")

output = ttk.Label(frame, text="")
output.pack(pady=10)

def generate():
    topic = topic_entry.get()               # get textbox value
    qtype = question_type.get()            # get selected radio button

    print("Topic:", topic)
    print("Type:", qtype)

    question_type_and_topic="{\"topic\":\"" + topic + "\", \"type\":\"" + qtype + "\"}" #wtf is this? I don't know, but it works. Don't touch it.

    client=genai.Client()

    print("generating prompt")

    with open(r"files\prompt", "r", encoding="utf-8") as f:
        prompt=f.read()

    print("sending API call")

    result=client.models.generate_content(
        model="gemini-3.1-flash-lite",
        contents=prompt+question_type_and_topic
    ).text

    print("writing json data to file")

    with open(r"files\result.json", "w", encoding="utf-8") as f:
        f.write(result)

    print("writing data to docx")

    def replace_in_paragraph(paragraph, old, new):
        #Replace text while preserving the formatting of the first run.
        if old not in paragraph.text:
            return
        # Merge all runs' text, then rewrite into the first run
        full_text = "".join(run.text for run in paragraph.runs)
        full_text = full_text.replace(old, new)
        
        for i, run in enumerate(paragraph.runs):
            run.text = full_text if i == 0 else ""

    def find_and_replace(filename, replacements: dict, output=None):
        doc = Document(filename)
        
        for paragraph in doc.paragraphs:
            for old, new in replacements.items():
                replace_in_paragraph(paragraph, old, new)
        
        # Also handle tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old, new in replacements.items():
                            replace_in_paragraph(paragraph, old, new)
        
        doc.save(output or filename)

    #to copy format docx (not needed since the function saves separately)

    #with open(r"files\Question.docx", "r", encoding="utf-8") as format:
    #    with open(q_fileName, "w", encoding="utf-8") as copy:
    #        copy.write(format.read())

    with open(r"files\result.json", "r", encoding="utf-8") as f:
        resdict=json.loads(f.read())

    filename=topic+".docx"

    if resdict["QType"]=="normal":
        find_and_replace(r"files\QuestionNormal.docx", {"Question English": resdict["QEng"], 
                                        "Question Sinhala": ConvertToLegacy(resdict["QSin"]),

                                        "Answer 1 English": resdict["AnswersEng"][0],
                                        "Answer 2 English": resdict["AnswersEng"][1],
                                        "Answer 3 English": resdict["AnswersEng"][2],
                                        "Answer 4 English": resdict["AnswersEng"][3],
                                        "Answer 5 English": resdict["AnswersEng"][4],
                                        "Answer 1 Sinhala": ConvertToLegacy(resdict["AnswersSin"][0]),
                                        "Answer 2 Sinhala": ConvertToLegacy(resdict["AnswersSin"][1]),
                                        "Answer 3 Sinhala": ConvertToLegacy(resdict["AnswersSin"][2]),
                                        "Answer 4 Sinhala": ConvertToLegacy(resdict["AnswersSin"][3]),
                                        "Answer 5 Sinhala": ConvertToLegacy(resdict["AnswersSin"][4]),
                                        
                                        "Explanation 1 English": resdict["ExplEng"][0],
                                        "Explanation 2 English": resdict["ExplEng"][1],
                                        "Explanation 3 English": resdict["ExplEng"][2],
                                        "Explanation 4 English": resdict["ExplEng"][3],
                                        "Explanation 5 English": resdict["ExplEng"][4],
                                        "Explanation 1 Sinhala": ConvertToLegacy(resdict["ExplSin"][0]),
                                        "Explanation 2 Sinhala": ConvertToLegacy(resdict["ExplSin"][1]),
                                        "Explanation 3 Sinhala": ConvertToLegacy(resdict["ExplSin"][2]),
                                        "Explanation 4 Sinhala": ConvertToLegacy(resdict["ExplSin"][3]),
                                        "Explanation 5 Sinhala": ConvertToLegacy(resdict["ExplSin"][4]),
                                        
                                        "QNum": str(resdict["AnsNo"]),
                                        },
                        output=filename)
        
    elif resdict["QType"]=="statement":
        find_and_replace(r"files\QuestionStatement.docx", {"Question English": resdict["QEng"], 
                                        "Question Sinhala": ConvertToLegacy(resdict["QSin"]),
                                        
                                        "StateAEng": resdict["StatementsEng"][0],
                                        "StateBEng": resdict["StatementsEng"][1],
                                        "StateCEng": resdict["StatementsEng"][2],
                                        "StateASin": ConvertToLegacy(resdict["StatementsSin"][0]),
                                        "StateBSin": ConvertToLegacy(resdict["StatementsSin"][1]),
                                        "StateCSin": ConvertToLegacy(resdict["StatementsSin"][2]),

                                        "Answer 1 English": resdict["AnswersEng"][0],
                                        "Answer 2 English": resdict["AnswersEng"][1],
                                        "Answer 3 English": resdict["AnswersEng"][2],
                                        "Answer 4 English": resdict["AnswersEng"][3],
                                        "Answer 5 English": resdict["AnswersEng"][4],

                                        "Answer 1 Sinhala": ConvertToLegacy(resdict["AnswersSin"][0]),
                                        "Answer 2 Sinhala": ConvertToLegacy(resdict["AnswersSin"][1]),
                                        "Answer 3 Sinhala": ConvertToLegacy(resdict["AnswersSin"][2]),
                                        "Answer 4 Sinhala": ConvertToLegacy(resdict["AnswersSin"][3]),
                                        "Answer 5 Sinhala": ConvertToLegacy(resdict["AnswersSin"][4]),
                                        
                                        "ExplAEng": resdict["ExplEng"][0],
                                        "ExplBEng": resdict["ExplEng"][1],
                                        "ExplCEng": resdict["ExplEng"][2],
                                        "ExplASin": ConvertToLegacy(resdict["ExplSin"][0]),
                                        "ExplBSin": ConvertToLegacy(resdict["ExplSin"][1]),
                                        "ExplCSin": ConvertToLegacy(resdict["ExplSin"][2]),
                                        
                                        "QNum": str(resdict["AnsNo"]),
                                        },
                        output=filename)
        
    elif resdict["QType"]=="code":
        find_and_replace(r"files\QuestionCode.docx", {
                                        "Question English": resdict["QEng"], 
                                        "Question Sinhala": ConvertToLegacy(resdict["QSin"]),

                                        "codelines": resdict["Code"],

                                        "Answer 1 English": resdict["AnswersEng"][0],
                                        "Answer 2 English": resdict["AnswersEng"][1],
                                        "Answer 3 English": resdict["AnswersEng"][2],
                                        "Answer 4 English": resdict["AnswersEng"][3],
                                        "Answer 5 English": resdict["AnswersEng"][4],
                                        "Answer 1 Sinhala": ConvertToLegacy(resdict["AnswersSin"][0]),
                                        "Answer 2 Sinhala": ConvertToLegacy(resdict["AnswersSin"][1]),
                                        "Answer 3 Sinhala": ConvertToLegacy(resdict["AnswersSin"][2]),
                                        "Answer 4 Sinhala": ConvertToLegacy(resdict["AnswersSin"][3]),
                                        "Answer 5 Sinhala": ConvertToLegacy(resdict["AnswersSin"][4]),
                                        
                                        "ExplanationEnglish": resdict["ExplEng"],
                                        "ExplanationSinhala": ConvertToLegacy(resdict["ExplSin"]),
                                        
                                        "QNum": str(resdict["AnsNo"]),
                                        },
                        output=filename)
        
    else:
        print("Unknown question type", resdict["QType"])

btn = ttk.Button(frame, text="Generate", command=generate)
btn.pack(pady=10)

root.mainloop()
